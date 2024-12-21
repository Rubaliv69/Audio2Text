import os
import logging
import subprocess
import tempfile
import wave
from pathlib import Path
from pydub import AudioSegment
import speech_recognition as sr
from docx import Document
from concurrent.futures import ThreadPoolExecutor, as_completed
import time
from typing import List, Tuple
from PyQt6.QtCore import QObject, pyqtSignal, Qt
import datetime
import gc
import sys

class AudioConverter(QObject):
    # Signaux pour la progression
    progress_updated = pyqtSignal(int, int)  # (segments_traités, total_segments)
    segment_completed = pyqtSignal(str)  # message de log pour chaque segment
    error_occurred = pyqtSignal(str)  # Signal pour les erreurs
    
    def __init__(self, max_workers=None):
        super().__init__()
        # Utiliser le nombre de threads CPU disponibles - 1 (minimum 1)
        self.max_workers = max_workers or max(1, os.cpu_count() - 1)
        self.language = 'fr-FR'  # Langue par défaut
        self.supported_formats = ['.wav', '.mp3', '.m4a', '.flac', '.ogg']
        self.recognizer = sr.Recognizer()
        self.is_running = False
        logging.info(f"Initialisation du convertisseur audio avec {self.max_workers} workers")

    def process_segment(self, segment_data):
        """Traite un segment audio et retourne le texte transcrit."""
        try:
            segment, segment_index, start_time, end_time = segment_data
            logging.debug(f"Segment {segment_index}/16: Tentative 1 de reconnaissance")
            
            # Utiliser un fichier temporaire pour le segment
            with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_file:
                segment_path = temp_file.name
                segment.export(segment_path, format="wav")
            
            # Créer un nouvel objet Recognizer pour chaque segment
            recognizer = sr.Recognizer()
            with sr.AudioFile(segment_path) as source:
                # Réduire l'utilisation de la mémoire en limitant la taille du buffer
                audio = recognizer.record(source, duration=min(45, end_time - start_time))
                text = recognizer.recognize_google(audio, language=self.language)
                
            # Supprimer le fichier temporaire
            try:
                os.unlink(segment_path)
                logging.debug(f"Segment {segment_index} supprimé")
            except Exception as e:
                logging.error(f"Erreur lors de la suppression du segment {segment_index}: {str(e)}")
            
            logging.debug(f"Segment {segment_index}/16: Reconnaissance réussie ({len(text)} caractères)")
            return segment_index, text.strip()
            
        except sr.UnknownValueError:
            logging.error(f"Segment {segment_index}/16: Audio incompréhensible")
            return segment_index, ""
            
        except sr.RequestError as e:
            logging.error(f"Segment {segment_index}/16: Erreur API ({str(e)})")
            return segment_index, ""
            
        except Exception as e:
            logging.error(f"Segment {segment_index}/16: Erreur inattendue ({str(e)})")
            return segment_index, ""
            
        finally:
            # S'assurer que le segment est libéré de la mémoire
            del segment
            gc.collect()

    def format_text(self, text):
        """Formate le texte pour une meilleure lisibilité"""
        # Ajouter une majuscule au début
        text = text.capitalize()
        
        # Ajouter un point à la fin si nécessaire
        if text and not text.endswith(('.', '!', '?')):
            text += '.'
        
        return text

    def convert_to_wav(self, audio_path):
        """Convertit le fichier audio en WAV"""
        # Créer un fichier temporaire avec extension .wav
        with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_file:
            wav_path = temp_file.name
        
        try:
            # Chercher ffmpeg dans le PATH
            ffmpeg_cmd = 'ffmpeg'
            if os.path.exists('/opt/homebrew/bin/ffmpeg'):
                ffmpeg_cmd = '/opt/homebrew/bin/ffmpeg'
            elif os.path.exists('/usr/local/bin/ffmpeg'):
                ffmpeg_cmd = '/usr/local/bin/ffmpeg'
            
            # Construire la commande ffmpeg
            command = [ffmpeg_cmd, '-i', audio_path, '-acodec', 'pcm_s16le', '-ac', '1', '-ar', '44100', '-y', wav_path]
            logging.info(f"Conversion en WAV: {' '.join(command)}")
            
            # Exécuter la commande
            result = subprocess.run(command, capture_output=True, text=True)
            if result.returncode != 0:
                raise Exception(f"Erreur ffmpeg: {result.stderr}")
            
            return wav_path
            
        except Exception as e:
            logging.error(f"Erreur lors de la conversion en WAV: {str(e)}")
            if os.path.exists(wav_path):
                os.unlink(wav_path)
            raise

    def get_audio_duration(self, wav_path):
        """Obtient la durée du fichier audio en secondes"""
        try:
            with wave.open(wav_path, 'rb') as wav_file:
                frames = wav_file.getnframes()
                rate = wav_file.getframerate()
                duration = frames / float(rate)
                logging.info(f"Durée du fichier: {duration:.1f} secondes")
                return duration
        except Exception as e:
            error_msg = f"Erreur lors de la lecture de la durée: {str(e)}"
            logging.error(error_msg)
            self.error_occurred.emit(error_msg)
            raise

    def split_audio(self, audio):
        """Divise le fichier audio en segments"""
        try:
            duration_ms = len(audio)
            segment_duration_ms = 45 * 1000
            segments = []
            
            for start_ms in range(0, duration_ms, segment_duration_ms):
                end_ms = min(start_ms + segment_duration_ms, duration_ms)
                
                # Extraire le segment
                segment = audio[start_ms:end_ms]
                
                # Calculer les temps en secondes
                start_s = start_ms / 1000
                end_s = end_ms / 1000
                
                # Stocker le segment avec ses informations
                segments.append((segment, len(segments) + 1, start_s, end_s))
                
                logging.info(f"Segment créé: {start_s:.1f}s - {end_s:.1f}s")
                
            return segments
            
        except Exception as e:
            error_msg = f"Erreur lors de la division de l'audio: {str(e)}"
            logging.error(error_msg)
            raise

    def save_to_word(self, text, output_path):
        """Sauvegarde le texte dans un document Word"""
        try:
            doc = Document()
            
            # Ajouter un titre
            doc.add_heading('Transcription Audio', 0)
            
            # Ajouter la date et l'heure
            now = datetime.datetime.now()
            doc.add_paragraph(f'Généré le {now.strftime("%d/%m/%Y à %H:%M")}')
            
            # Ajouter une ligne de séparation
            doc.add_paragraph('_' * 50)
            
            # Ajouter le texte transcrit
            doc.add_paragraph(text)
            
            # Sauvegarder le document
            doc.save(output_path)
            logging.info(f"Document Word sauvegardé: {output_path}")
            
        except Exception as e:
            error_msg = f"Erreur lors de la sauvegarde du document Word: {str(e)}"
            logging.error(error_msg)
            self.error_occurred.emit(error_msg)
            raise

    def convert_to_text(self, audio_path: str, language: str = None) -> str:
        """Convertit un fichier audio en texte"""
        try:
            self.is_running = True
            if language:
                self.language = language
            
            logging.info(f"Début de la conversion de {audio_path} en texte (langue: {self.language})")
            
            # Vérifier si le fichier existe
            if not os.path.exists(audio_path):
                raise FileNotFoundError(f"Le fichier {audio_path} n'existe pas")
            
            # Convertir en WAV si nécessaire
            wav_path = self.convert_to_wav(audio_path)
            logging.info(f"Fichier converti en WAV : {wav_path}")
            
            # Charger l'audio
            audio = AudioSegment.from_wav(wav_path)
            logging.info(f"Fichier audio chargé, durée : {len(audio)/1000} secondes")
            
            # Diviser en segments
            segments = self.split_audio(audio)
            total_segments = len(segments)
            logging.info(f"Audio divisé en {total_segments} segments")
            
            # Initialiser le résultat
            result_text = []
            segments_processed = 0
            
            # Créer un pool de threads
            with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                # Soumettre les tâches
                futures = []
                for segment in segments:
                    if not self.is_running:
                        logging.info("Conversion interrompue")
                        break
                    futures.append(executor.submit(self.process_segment, segment))
                
                # Traiter les résultats
                for future in futures:
                    if not self.is_running:
                        break
                    try:
                        index, text = future.result()
                        result_text.append(text)
                        segments_processed += 1
                        progress = (segments_processed * 100) // total_segments
                        self.progress_updated.emit(segments_processed, total_segments)
                        self.segment_completed.emit(f"Segment {index} traité")
                    except Exception as e:
                        error_msg = f"Erreur lors du traitement d'un segment : {str(e)}"
                        logging.error(error_msg)
                        self.error_occurred.emit(error_msg)
            
            # Nettoyer le fichier WAV temporaire
            if wav_path != audio_path:
                try:
                    os.unlink(wav_path)
                    logging.info("Fichier WAV temporaire supprimé")
                except Exception as e:
                    logging.warning(f"Impossible de supprimer le fichier temporaire : {str(e)}")
            
            # Formater et retourner le texte final
            final_text = " ".join(result_text)
            final_text = self.format_text(final_text)
            logging.info("Conversion terminée avec succès")
            return final_text
            
        except Exception as e:
            error_msg = f"Erreur lors de la conversion : {str(e)}"
            logging.error(error_msg)
            self.error_occurred.emit(error_msg)
            raise
        finally:
            self.is_running = False

    def convert_audio(self, input_file: str, output_format: str = 'wav') -> str:
        """Convertit un fichier audio dans le format spécifié."""
        input_ext = os.path.splitext(input_file)[1].lower()
        if input_ext not in self.supported_formats:
            raise ValueError(f"Format non supporté : {input_ext}")

        if not os.path.exists(input_file):
            raise FileNotFoundError(f"Le fichier {input_file} n'existe pas")

        # Le reste de votre code de conversion existant
        return input_file  # Pour le moment, retourne simplement le fichier d'entrée

    def get_duration(self, file_path: str) -> float:
        """Obtient la durée d'un fichier audio en secondes."""
        try:
            logging.info(f"Obtention de la durée pour le fichier : {file_path}")
            if not os.path.exists(file_path):
                error_msg = f"Le fichier {file_path} n'existe pas"
                logging.error(error_msg)
                raise FileNotFoundError(error_msg)
            
            # Vérifier l'extension du fichier
            ext = os.path.splitext(file_path)[1].lower()
            logging.info(f"Extension du fichier : {ext}")
            if ext not in self.supported_formats:
                error_msg = f"Format de fichier non supporté : {ext}"
                logging.error(error_msg)
                raise ValueError(error_msg)
            
            # Charger le fichier audio
            try:
                audio = AudioSegment.from_file(file_path)
                duration = len(audio) / 1000.0  # Convertir en secondes
                logging.info(f"Durée obtenue : {duration} secondes")
                return duration
            except Exception as e:
                error_msg = f"Erreur lors du chargement du fichier audio : {str(e)}"
                logging.error(error_msg)
                raise Exception(error_msg)
                
        except Exception as e:
            logging.error(f"Erreur dans get_duration : {str(e)}")
            raise
